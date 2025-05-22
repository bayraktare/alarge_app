import sqlite3, qrcode, uuid
import os
from docx import Document
from fpdf import FPDF
from reportlab.pdfgen import canvas
from reportlab.platypus import (SimpleDocTemplate, Paragraph, PageBreak, Image, Spacer, Table, TableStyle)
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.pagesizes import inch, A4
from reportlab.graphics.shapes import Drawing, Line
from reportlab.lib.colors import Color
from io import BytesIO
from pdf import PDFPSReporte
from modules.customWidgets import *


from PySide6.QtGui import QFont
from PySide6.QtCore import Qt
from PySide6.QtWidgets import (QInputDialog, QGridLayout, QPushButton, QComboBox,
                             QRadioButton, QLabel, QLineEdit, QVBoxLayout, QWidget,
                             QHBoxLayout, QMessageBox, QFileDialog)




HOME = os.getcwd()

class CompanyInfoWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Enter Company Information")
        self.setGeometry(100, 100, 600, 400)

        self.initUI()

    def initUI(self):
        # Create labels and input fields for company information
        self.lbl_company_name = QLabel("Company Name:", self)
        self.lbl_logo = QLabel("Logo:", self)
        self.lbl_phone = QLabel("Phone:", self)
        self.lbl_email = QLabel("Email:", self)
        self.lbl_address = QLabel("Address:", self)
        self.lbl_operator_names = QLabel("Operator Name(s):", self)
        self.lbl_machine_names = QLabel("Machine Names:", self)
        self.lbl_test_names = QLabel("Test Names:", self)
        self.lbl_ip_addresses = QLabel("IP Addresses:", self)

        self.input_company_name = QLineEdit(self)
        self.input_logo = QLineEdit(self)
        self.input_phone = PhoneNumLineEdit(self)
        self.input_email = QLineEdit(self)
        self.input_address = QLineEdit(self)

        self.operator_layout = QVBoxLayout()
        self.add_operator_row()
        self.machine_layout = QVBoxLayout()
        self.add_machine_row()
        self.test_layout = QVBoxLayout()
        self.add_test_row()
        self.ip_layout = QVBoxLayout()
        self.add_ip_row()

        self.btn_save = QPushButton("Save", self)
        self.btn_save.clicked.connect(self.saveCompanyInfo)

        # Set font
        font = QFont()
        font.setPointSize(12)
        font.setBold(True)

        self.lbl_company_name.setFont(font)
        self.lbl_logo.setFont(font)
        self.lbl_phone.setFont(font)
        self.lbl_email.setFont(font)
        self.lbl_address.setFont(font)
        self.lbl_operator_names.setFont(font)
        self.lbl_machine_names.setFont(font)
        self.lbl_test_names.setFont(font)
        self.lbl_ip_addresses.setFont(font)

        # Set styles
        self.setStyleSheet("""
            QLabel {
                color: rgb(221, 221, 221); /* Updated text color */
                font-weight: bold;
            }
            QLineEdit {
                border: 2px solid rgb(44, 49, 58); /* Updated border color */
                border-radius: 5px;
                padding: 5px;
                background-color: rgb(33, 37, 43); /* Updated background color */
                color: rgb(221, 221, 221); /* Updated text color */
            }
            QPushButton {
                background-color: rgb(136, 36, 40); /* Updated button color */
                color: rgb(255, 255, 255); /* Updated text color */
                border-radius: 5px;
                padding: 8px 16px;
            }
            QPushButton:hover {
                background-color: rgb(189, 147, 249); /* Updated hover color */
            }
        """)

        # Create layout
        grid_layout = QVBoxLayout()

        # Company info labels and inputs
        company_info_layout = QVBoxLayout()
        company_info_layout.addWidget(self.lbl_company_name)
        company_info_layout.addWidget(self.input_company_name)
        company_info_layout.addWidget(self.lbl_logo)
        company_info_layout.addWidget(self.input_logo)
        company_info_layout.addWidget(self.lbl_phone)
        company_info_layout.addWidget(self.input_phone)
        company_info_layout.addWidget(self.lbl_email)
        company_info_layout.addWidget(self.input_email)
        company_info_layout.addWidget(self.lbl_address)
        company_info_layout.addWidget(self.input_address)
        grid_layout.addLayout(company_info_layout)

        # Operator names
        operator_layout = QVBoxLayout()
        operator_layout.addWidget(self.lbl_operator_names)
        operator_layout.addLayout(self.operator_layout)

        buttons_layout = QHBoxLayout()
        self.btn_add_operator = QPushButton("+ Add Operator", self)
        self.btn_add_operator.clicked.connect(self.add_operator_row)
        self.btn_remove_operator = QPushButton("- Remove Operator", self)
        self.btn_remove_operator.clicked.connect(self.remove_operator_row)
        buttons_layout.addWidget(self.btn_add_operator)
        buttons_layout.addWidget(self.btn_remove_operator)

        operator_layout.addLayout(buttons_layout)
        grid_layout.addLayout(operator_layout)

        # Machine names
        machine_layout = QVBoxLayout()
        machine_layout.addWidget(self.lbl_machine_names)
        machine_layout.addLayout(self.machine_layout)

        buttons_layout = QHBoxLayout()
        self.btn_add_machine = QPushButton("+ Add Machine", self)
        self.btn_add_machine.clicked.connect(self.add_machine_row)
        self.btn_remove_machine = QPushButton("- Remove Machine", self)
        self.btn_remove_machine.clicked.connect(self.remove_machine_row)
        buttons_layout.addWidget(self.btn_add_machine)
        buttons_layout.addWidget(self.btn_remove_machine)

        machine_layout.addLayout(buttons_layout)
        grid_layout.addLayout(machine_layout)

        # Test names
        test_layout = QVBoxLayout()
        test_layout.addWidget(self.lbl_test_names)
        test_layout.addLayout(self.test_layout)

        buttons_layout = QHBoxLayout()
        self.btn_add_test = QPushButton("+ Add Test", self)
        self.btn_add_test.clicked.connect(self.add_test_row)
        self.btn_remove_test = QPushButton("- Remove Test", self)
        self.btn_remove_test.clicked.connect(self.remove_test_row)
        buttons_layout.addWidget(self.btn_add_test)
        buttons_layout.addWidget(self.btn_remove_test)

        test_layout.addLayout(buttons_layout)
        grid_layout.addLayout(test_layout)

        # IP addresses
        ip_layout = QVBoxLayout()
        ip_layout.addWidget(self.lbl_ip_addresses)
        ip_layout.addLayout(self.ip_layout)

        buttons_layout = QHBoxLayout()
        self.btn_add_ip = QPushButton("+ Add IP Address", self)
        self.btn_add_ip.clicked.connect(self.add_ip_row)
        self.btn_remove_ip = QPushButton("- Remove IP Address", self)
        self.btn_remove_ip.clicked.connect(self.remove_ip_row)
        buttons_layout.addWidget(self.btn_add_ip)
        buttons_layout.addWidget(self.btn_remove_ip)

        ip_layout.addLayout(buttons_layout)
        grid_layout.addLayout(ip_layout)

        # Save button
        grid_layout.addWidget(self.btn_save, alignment=Qt.AlignHCenter)

        self.setLayout(grid_layout)

    def add_operator_row(self):
        self.add_row(self.operator_layout, "Operator Name")

    def remove_operator_row(self):
        self.remove_row(self.operator_layout)

    def add_machine_row(self):
        self.add_row(self.machine_layout, "Machine Name")

    def remove_machine_row(self):
        self.remove_row(self.machine_layout)

    def add_test_row(self):
        self.add_row(self.test_layout, "Test Name")

    def remove_test_row(self):
        self.remove_row(self.test_layout)

    def add_ip_row(self):
        hbox = QHBoxLayout()
        input_machine_name = QLineEdit(self)
        input_machine_name.setPlaceholderText("Machine Name")
        input_ip_address = IPAddressLineEdit(self)
        input_ip_address.setPlaceholderText("IP Address")
        hbox.addWidget(input_machine_name)
        hbox.addWidget(input_ip_address)
        self.ip_layout.addLayout(hbox)

    def remove_ip_row(self):
        if self.ip_layout.count() > 0:
            layout_item = self.ip_layout.takeAt(self.ip_layout.count() - 1)
            for i in reversed(range(layout_item.count())):
                layout_item.itemAt(i).widget().setParent(None)

    def add_row(self, layout, placeholder):
        hbox = QHBoxLayout()
        input_field = QLineEdit(self)
        input_field.setPlaceholderText(placeholder)
        hbox.addWidget(input_field)
        layout.addLayout(hbox)

    def remove_row(self, layout):
        if layout.count() > 0:
            layout_item = layout.takeAt(layout.count() - 1)
            for i in reversed(range(layout_item.count())):
                layout_item.itemAt(i).widget().setParent(None)

    # Function to save entered company information to a text file
    def saveCompanyInfo(self):
        # Prompt user to specify file name and path
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Company Information", "", "Text Files (*.txt)")
        if file_path:
            with open(file_path, 'w') as file:
                file.write("Company Name: {}\n".format(self.input_company_name.text()))
                file.write("Logo: {}\n".format(self.input_logo.text()))
                file.write("Phone: {}\n".format(self.input_phone.text()))
                file.write("Email: {}\n".format(self.input_email.text()))
                file.write("Address: {}\n".format(self.input_address.text()))
                file.write("Operator Name(s): {}\n".format(self.lbl_operator_names.text()))
                file.write("Machine Names: {}\n".format(self.lbl_machine_names.text()))
                file.write("Test Names: {}\n".format(self.lbl_test_names.text()))
                file.write("IP Addresses: {}\n".format(self.lbl_ip_addresses.text()))

class NewRecordWindow(QWidget):
    
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Material Description")
        self.setGeometry(100, 100, 600, 400)

        self.initUI()

    def initUI(self):
        # Create labels and input fields for material description
        self.lbl_raw_material = QLabel("Raw Material Information (e.g., H2O):", self)
        self.input_raw_material = QLineEdit(self)
        self.input_raw_material.setPlaceholderText("Chemical Formula")

        self.lbl_component_ratio = QLabel("Component Ratio:", self)
        self.component_layout = QVBoxLayout()
        self.add_component_row()

        self.btn_add_component = QPushButton("+ Add Component", self)
        self.btn_add_component.clicked.connect(self.add_component_row)
        self.btn_remove_component = QPushButton("- Remove Component", self)
        self.btn_remove_component.clicked.connect(self.remove_component_row)

        # self.lbl_supplier = QLabel("Supplier:", self)
        # self.input_supplier = QLineEdit(self)

        self.lbl_manufacturing_date = QLabel("Manufacturing Date:", self)
        self.input_manufacturing_date = DateLineEdit(self)

        self.btn_save_sql = QPushButton("Save to SQL Database", self)
        self.btn_save_sql.clicked.connect(self.save_to_sql)

        self.btn_save_txt = QPushButton("Save to Text File", self)
        self.btn_save_txt.clicked.connect(self.save_to_txt)

        self.btn_generate_qr = QPushButton("Generate QR Code", self)
        self.btn_generate_qr.clicked.connect(self.make_qrcode)

        self.btn_generate_pdf = QPushButton("Generate PDF", self)
        self.btn_generate_pdf.clicked.connect(self.make_pdf)

        # Set font
        font = QFont()
        font.setPointSize(12)
        font.setBold(True)

        self.lbl_raw_material.setFont(font)
        self.lbl_component_ratio.setFont(font)
        # self.lbl_supplier.setFont(font)
        self.lbl_manufacturing_date.setFont(font)

        # Set styles
        self.setStyleSheet("""
            QLabel {
                color: rgb(221, 221, 221); /* Updated text color */
                font-weight: bold;
            }
            QLineEdit {
                border: 2px solid rgb(44, 49, 58); /* Updated border color */
                border-radius: 5px;
                padding: 5px;
                background-color: rgb(33, 37, 43); /* Updated background color */
                color: rgb(221, 221, 221); /* Updated text color */
            }
            QPushButton {
                background-color: rgb(136, 36, 40); /* Updated button color */
                color: rgb(255, 255, 255); /* Updated text color */
                border-radius: 5px;
                padding: 8px 16px;
            }
            QPushButton:hover {
                background-color: rgb(189, 147, 249); /* Updated hover color */
            }
        """)

        # Create layout
        layout = QVBoxLayout()
        layout.addWidget(self.lbl_raw_material)
        layout.addWidget(self.input_raw_material)

        layout.addWidget(self.lbl_component_ratio)
        layout.addLayout(self.component_layout)

        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.btn_add_component)
        buttons_layout.addWidget(self.btn_remove_component)
        layout.addLayout(buttons_layout)

        # layout.addWidget(self.lbl_supplier)
        # layout.addWidget(self.input_supplier)
        layout.addWidget(self.lbl_manufacturing_date)
        layout.addWidget(self.input_manufacturing_date)

        layout.addWidget(self.btn_save_sql)
        layout.addWidget(self.btn_save_txt)
        layout.addWidget(self.btn_generate_qr)
        layout.addWidget(self.btn_generate_pdf)

        self.setLayout(layout)

    def add_component_row(self):
        hbox = QHBoxLayout()
        input_name = QLineEdit(self)
        input_name.setPlaceholderText("Component Name")
        input_percent = PercentageLineEdit(self)
        supplier_name = QLineEdit(self)
        supplier_name.setPlaceholderText("Supplier Name")
        hbox.addWidget(input_name)
        hbox.addWidget(input_percent)
        hbox.addWidget(supplier_name)
        self.component_layout.addLayout(hbox)

    def remove_component_row(self):
        if self.component_layout.count() > 0:
            layout_item = self.component_layout.takeAt(self.component_layout.count() - 1)
            for i in reversed(range(layout_item.count())):
                layout_item.itemAt(i).widget().setParent(None)

    def save_to_sql(self):
        # Create database connection
        conn = sqlite3.connect('material_records.db')
        cursor = conn.cursor()

        # Create table if it doesn't exist
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS materials (
                id INTEGER PRIMARY KEY,
                raw_material TEXT,
                supplier TEXT,
                manufacturing_date TEXT
            )
        ''')

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS components (
                material_id INTEGER,
                component_name TEXT,
                component_percent TEXT,
                FOREIGN KEY(material_id) REFERENCES materials(id)
            )
        ''')

        # Insert material information
        cursor.execute('''
            INSERT INTO materials (raw_material, supplier, manufacturing_date)
            VALUES (?, ?, ?)
        ''', (self.input_raw_material.text(), self.input_supplier.text(), self.input_manufacturing_date.text()))

        material_id = cursor.lastrowid

        # Insert component information
        for i in range(self.component_layout.count()):
            layout_item = self.component_layout.itemAt(i)
            component_name = layout_item.itemAt(0).widget().text()
            component_percent = layout_item.itemAt(1).widget().text()
            cursor.execute('''
                INSERT INTO components (material_id, component_name, component_percent)
                VALUES (?, ?, ?)
            ''', (material_id, component_name, component_percent))

        conn.commit()
        conn.close()

        # Inform the user
        QMessageBox.information(self, "Success", "Record saved to SQL database successfully.")

    def save_to_txt(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save to Text File", "", "Text Files (*.txt);;All Files (*)", options=options)
        if file_path:
            with open(file_path, 'w') as file:
                file.write(f"Raw Material: {self.input_raw_material.text()}\n")
                file.write("Components:\n")
                for i in range(self.component_layout.count()):
                    layout_item = self.component_layout.itemAt(i)
                    component_name = layout_item.itemAt(0).widget().text()
                    component_percent = layout_item.itemAt(1).widget().text()
                    supplier_name = layout_item.itemAt(2).widget().text()
                    file.write(f"  - {component_name}: {component_percent}% {supplier_name}\n")
                # file.write(f"Supplier: {self.input_supplier.text()}\n")
                file.write(f"Manufacturing Date: {self.input_manufacturing_date.text()}\n")

            # Inform the user
            QMessageBox.information(self, "Success", "Record saved to text file successfully.")

    def make_qrcode(self):

        
        if not self.input_raw_material.text() or not self.input_manufacturing_date.text():
            QMessageBox.warning(self, "Input Error", "Please fill in all the fields before generating a QR code.")
            return
        
        else:
            options = QFileDialog.Options()
            file_path, _ = QFileDialog.getSaveFileName(self, "Save QR Code", "", "PNG Files (*.png);;All Files (*)", options=options)
            if file_path:
                text = f"Raw Material: {self.input_raw_material.text()}\nComponents:\n"
                for i in range(self.component_layout.count()):
                    layout_item = self.component_layout.itemAt(i)
                    component_name = layout_item.itemAt(0).widget().text()
                    component_percent = layout_item.itemAt(1).widget().text()
                    supplier_name = layout_item.itemAt(2).widget().text()
                    text += f"  - {component_name}: {component_percent}% {supplier_name}\n"
                text += f"Manufacturing Date: {self.input_manufacturing_date.text()}"
                qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
                qr.add_data(text)
                qr.make(fit=True)
                img = qr.make_image(fill_color="black", back_color="white")
                img.save(file_path)
            
        QMessageBox.information(self, "Success", "QR Code generated and saved successfully.")

    



    def make_pdf(self):

    
        # if not self.input_raw_material.text() or not self.input_manufacturing_date.text():
        #     QMessageBox.warning(self, "Input Error", "Please fill in all the fields before generating a QR code.")
        #     return
        # else:
            PDFPSReporte('deneme.pdf',input_raw_material=self.input_raw_material, input_manufacturing_date=self.input_manufacturing_date, component_layout=self.component_layout)

            QMessageBox.information(self, "Success", "PDF generated and saved successfully.")

        

        

        


        

            

    # def generate_qr_code(self): #eski qr code
    #     # if not self.input_raw_material.text() or not self.input_supplier.text() or not self.input_manufacturing_date.text():
    #     if not self.input_raw_material.text() or not self.input_manufacturing_date.text():
    #         QMessageBox.warning(self, "Input Error", "Please fill in all the fields before generating a QR code.")
    #         return

    #     # data = f"Raw Material: {self.input_raw_material.text()}, Supplier: {self.input_supplier.text()}, Manufacturing Date: {self.input_manufacturing_date.text()}, Components: "
    #     data = f"Raw Material: {self.input_raw_material.text()}, Manufacturing Date: {self.input_manufacturing_date.text()}, Components: "
    #     components = []
    #     for i in range(self.component_layout.count()):
    #         layout_item = self.component_layout.itemAt(i)
    #         component_name = layout_item.itemAt(0).widget().text()
    #         component_percent = layout_item.itemAt(1).widget().text()
    #         supplier_name = layout_item.itemAt(2).widget().text()
    #         components.append(f"{component_name} ({component_percent}%) {supplier_name}")
    #     data += ", ".join(components)

    #     qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
    #     qr.add_data(data)
    #     qr.make(fit=True)

    #     img = qr.make_image(fill_color="black", back_color="white")

    #     options = QFileDialog.Options()
    #     file_path, _ = QFileDialog.getSaveFileName(self, "Save QR Code", "", "PNG Files (*.png);;All Files (*)", options=options)
    #     if file_path:
    #         img.save(file_path)
    #         QMessageBox.information(self, "Success", "QR Code generated and saved successfully.")

class ExistingRecordWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.successfullyBuilt = False
        self.setWindowTitle("Existing Record")
        self.setGeometry(100, 100, 600, 400)

        self.initUI()

    def initUI(self):
        # Check if abc.txt exists
        if not self.read_company_info():
            QMessageBox.warning(self, "Warning", "Company information file (abc.txt) not found!")
            self.close()
            return
        
        # Let the user choose save records
        if self.choose_saved_record():
            # Let the user choose another file about machines and tests
            if self.choose_machine_test_file():
                # Create UI components for saving the merged data
                self.create_save_buttons()
                self.successfullyBuilt = True

    def read_company_info(self):
        try:
            with open("abc.txt", "r") as file:
                lines = file.readlines()
                self.company_info = {}
                for line in lines:
                    key, value = line.strip().split(":", 1)
                    self.company_info[key.strip()] = value.strip()
            return True
        except FileNotFoundError:
            return False

    def choose_saved_record(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(self, "Choose Saved Record", "", "Database Files (*.db);;Text Files (*.txt);;QR Code Files (*.qrcode);;All Files (*)", options=options)
        if file_path.endswith('txt'):
            try:
                with open(file_path, 'r') as file:
                    self.saved_record_content = file.read()
                    return True
            except FileNotFoundError:
                return False
        elif file_path.endswith('.db'):
            return True
            pass
            
    def choose_machine_test_file(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(self, "Choose Machine and Test File", "", "Text Files (*.txt);;QR Code Files (*.qrcode);;All Files (*)", options=options)
        
        try:
            with open(file_path, 'r') as file:
                self.machine_test_content = file.read()
                return True
        except FileNotFoundError:
            return False

    def create_save_buttons(self):
        self.btn_save_txt = QPushButton("Save as Text File", self)
        self.btn_save_txt.clicked.connect(self.save_to_txt)

        self.btn_save_qr = QPushButton("Save as QR Code", self)
        self.btn_save_qr.clicked.connect(self.save_to_qrcode)

        self.btn_save_word = QPushButton("Save as Word Document", self)
        self.btn_save_word.clicked.connect(self.save_to_word)

        self.btn_save_pdf = QPushButton("Save as PDF", self)
        self.btn_save_pdf.clicked.connect(self.save_to_pdf)

        # Set font
        font = QFont()
        font.setPointSize(12)
        font.setBold(True)

        self.btn_save_txt.setFont(font)
        self.btn_save_qr.setFont(font)
        self.btn_save_word.setFont(font)
        self.btn_save_pdf.setFont(font)

        # Set styles
        self.setStyleSheet("""
            QPushButton {
                background-color: rgb(136, 36, 40); /* Updated button color */
                color: rgb(255, 255, 255); /* Updated text color */
                font-weight: bold;
                border-radius: 10px;
                padding: 20px 40px;
                border: 2px solid rgb(44, 49, 58); /* Updated border color */
            }
            QPushButton:hover {
                background-color: rgb(189, 147, 249); /* Updated hover color */
            }
            QLabel {
                color: rgb(221, 221, 221); /* Updated text color */
                font-weight: bold;
            }
            QComboBox {
                color: rgb(221, 221, 221); /* Updated text color */
                background-color: rgb(33, 37, 43); /* Updated background color */
                border: 2px solid rgb(44, 49, 58); /* Updated border color */
            }
        """)

        # Create layout
        grid_layout = QGridLayout()
        grid_layout.addWidget(self.btn_save_txt, 0, 0)
        grid_layout.addWidget(self.btn_save_qr, 0, 1)
        grid_layout.addWidget(self.btn_save_word, 1, 0)
        grid_layout.addWidget(self.btn_save_pdf, 1, 1)

        # Set layout to the widget
        layout = QVBoxLayout()
        layout.addLayout(grid_layout)
        self.setLayout(layout)

    def merge_contents(self):
        unique_identifier = uuid.uuid4()
        merged_content = f"Company Information:\n{self.company_info}\n\n"
        merged_content += f"Saved Record Content:\n{self.saved_record_content}\n\n"
        merged_content += f"Machine and Test Content:\n{self.machine_test_content}\n\n"
        merged_content += f"Unique ID: {unique_identifier}"
        return merged_content

    def save_to_txt(self):
        merged_content = self.merge_contents()
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as Text File", "", "Text Files (*.txt);;All Files (*)", options=options)
        if file_path:
            with open(file_path, 'w') as file:
                file.write(merged_content)
            QMessageBox.information(self, "Success", "Text file saved successfully.")

    def save_to_qrcode(self):
        merged_content = self.merge_contents()
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
        qr.add_data(merged_content)
        qr.make(fit=True)

        img = qr.make_image(fill_color="black", back_color="white")
        
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save QR Code", "", "PNG Files (*.png);;All Files (*)", options=options)
        if file_path:
            img.save(file_path)
            QMessageBox.information(self, "Success", "QR Code saved successfully.")

    def save_to_word(self):
        merged_content = self.merge_contents()
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as Word Document", "", "Word Files (*.docx);;All Files (*)", options=options)
        if file_path:
            doc = Document()
            doc.add_heading('Merged Report', 0)
            doc.add_paragraph(merged_content)
            doc.save(file_path)
            QMessageBox.information(self, "Success", "Word document saved successfully.")

    def save_to_pdf(self):
        merged_content = self.merge_contents()
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as PDF", "", "PDF Files (*.pdf);;All Files (*)", options=options)
        if file_path:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, merged_content)
            pdf.output(file_path)
            QMessageBox.information(self, "Success", "PDF saved successfully.")

class ExistingTestWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Existing Test")
        self.setGeometry(100, 100, 600, 400)

        self.initUI()

    def initUI(self):
        try:
            with open("abc.txt", "r") as file:
                lines = file.readlines()
                self.machine_names = []
                self.test_names = []
                for line in lines:
                    if line.startswith("Machine:"):
                        self.machine_names.append(line.strip().split(":")[1].strip())
                    elif line.startswith("Test:"):
                        self.test_names.append(line.strip().split(":")[1].strip())
        except FileNotFoundError:
            pass
        
        if len(self.machine_names) == 0:
            # Read "abc.txt" to get machine names and test names, use default values if file reading fails
            self.machine_names = ["Extensiometer", "Gas", "EB"]
            self.test_names = ["MFI", "Vikat", "Universal"]
        # Create combo box for machine names
        self.lbl_machine_names = QLabel("Machine Names:", self)
        self.combo_machine_names = QComboBox(self)
        self.combo_machine_names.addItems(self.machine_names)

        # Create layout for test names and count
        self.test_layout = QVBoxLayout()
        for test_name in self.test_names:
            hbox = QHBoxLayout()
            radio_button = QRadioButton(test_name, self)
            input_count = QLineEdit(self)
            input_count.setPlaceholderText("Count")
            hbox.addWidget(radio_button)
            hbox.addWidget(input_count)
            self.test_layout.addLayout(hbox)

        self.btn_save_txt = QPushButton("Save to Text File", self)
        self.btn_save_txt.clicked.connect(self.save_to_txt)

        self.btn_save_sql = QPushButton("Save to SQL Database", self)
        self.btn_save_sql.clicked.connect(self.save_to_sql)

        self.btn_add_new_test_same_sample = QPushButton("Add New Test to Same Sample", self)
        self.btn_add_new_test_same_sample.clicked.connect(self.add_new_test_same_sample)

        self.btn_add_new_test_different_sample = QPushButton("Add New Test for Different Sample", self)
        self.btn_add_new_test_different_sample.clicked.connect(self.add_new_test_different_sample)

        self.btn_generate_qr = QPushButton("Generate QR Code", self)
        self.btn_generate_qr.clicked.connect(self.generate_qr_code)

        # Set font
        font = QFont()
        font.setPointSize(12)
        font.setBold(True)

        self.lbl_machine_names.setFont(font)

        # Set styles
        self.setStyleSheet("""
            QLabel {
                color: rgb(221, 221, 221); /* Updated text color */
                font-weight: bold;
            }
            QLineEdit {
                border: 2px solid rgb(44, 49, 58); /* Updated border color */
                border-radius: 5px;
                padding: 5px;
                background-color: rgb(33, 37, 43); /* Updated background color */
                color: rgb(221, 221, 221); /* Updated text color */
            }
            QPushButton {
                background-color: rgb(136, 36, 40); /* Updated button color */
                color: rgb(255, 255, 255); /* Updated text color */
                border-radius: 5px;
                padding: 8px 16px;
            }
            QPushButton:hover {
                background-color: rgb(189, 147, 249); /* Updated hover color */
            }
        """)

        # Create layout
        layout = QVBoxLayout()
        layout.addWidget(self.lbl_machine_names)
        layout.addWidget(self.combo_machine_names)

        layout.addLayout(self.test_layout)
        layout.addWidget(self.btn_save_txt)
        layout.addWidget(self.btn_save_sql)
        layout.addWidget(self.btn_add_new_test_same_sample)
        layout.addWidget(self.btn_add_new_test_different_sample)
        layout.addWidget(self.btn_generate_qr)

        self.setLayout(layout)

    def save_to_txt(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save to Text File", "", "Text Files (*.txt);;All Files (*)", options=options)
        if file_path:
            with open(file_path, 'w') as file:
                file.write("Machine Name:\n")
                file.write(f"  - {self.combo_machine_names.currentText()}\n")

                file.write("Tests:\n")
                for i in range(self.test_layout.count()):
                    layout_item = self.test_layout.itemAt(i)
                    radio_button = layout_item.itemAt(0).widget()
                    input_count = layout_item.itemAt(1).widget()
                    if radio_button.isChecked():
                        file.write(f"  - {radio_button.text()}: {input_count.text()}\n")

            # Inform the user
            QMessageBox.information(self, "Success", "Record saved to text file successfully.")

    def save_to_sql(self):
        conn = sqlite3.connect('existing_records.db')
        cursor = conn.cursor()

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS records (
                id INTEGER PRIMARY KEY,
                machine_name TEXT,
                test_name TEXT,
                test_count INTEGER
            )
        ''')

        machine_name = self.combo_machine_names.currentText()
        for i in range(self.test_layout.count()):
            layout_item = self.test_layout.itemAt(i)
            radio_button = layout_item.itemAt(0).widget()
            input_count = layout_item.itemAt(1).widget()
            if radio_button.isChecked():
                cursor.execute('''
                    INSERT INTO records (machine_name, test_name, test_count)
                    VALUES (?, ?, ?)
                ''', (machine_name, radio_button.text(), input_count.text()))

        conn.commit()
        conn.close()

        QMessageBox.information(self, "Success", "Record saved to SQL database successfully.")

    def add_new_test_same_sample(self):
        self.add_test_widgets()

    def add_new_test_different_sample(self):
        machine_name, ok = QInputDialog.getText(self, "New Sample", "Enter new machine name:")
        if ok and machine_name:
            self.combo_machine_names.addItem(machine_name)
            self.combo_machine_names.setCurrentText(machine_name)
            self.add_test_widgets()

    def add_test_widgets(self):
        for test_name in self.test_names:
            hbox = QHBoxLayout()
            radio_button = QRadioButton(test_name, self)
            input_count = QLineEdit(self)
            input_count.setPlaceholderText("Count")
            hbox.addWidget(radio_button)
            hbox.addWidget(input_count)
            self.test_layout.addLayout(hbox)

    def generate_qr_code(self):
        if self.combo_machine_names.currentText() == "" or all(not self.test_layout.itemAt(i).itemAt(0).widget().isChecked() for i in range(self.test_layout.count())):
            QMessageBox.warning(self, "Input Error", "Please select a machine name and at least one test before generating a QR code.")
            return

        data = f"Machine Name: {self.combo_machine_names.currentText()}; Tests: "
        tests = []
        for i in range(self.test_layout.count()):
            layout_item = self.test_layout.itemAt(i)
            radio_button = layout_item.itemAt(0).widget()
            input_count = layout_item.itemAt(1).widget()
            if radio_button.isChecked():
                tests.append(f"{radio_button.text()} ({input_count.text()})")
        data += ", ".join(tests)

        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
        qr.add_data(data)
        qr.make(fit=True)

        img = qr.make_image(fill_color="black", back_color="white")

        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(self, "Save QR Code", "", "PNG Files (*.png);;All Files (*)", options=options)
        if file_path:
            img.save(file_path)
            QMessageBox.information(self, "Success", "QR Code generated and saved successfully.")

class GenerateReportWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Generate Report")
        self.setGeometry(100, 100, 600, 400)

        self.initUI()

    def initUI(self):

        try:
            with open("abc.txt", "r") as file:
                lines = file.readlines()
                self.machine_names = []
                self.test_names = []
                for line in lines:
                    if line.startswith("Machine:"):
                        self.machine_names.append(line.strip().split(":")[1].strip())
                    elif line.startswith("Test:"):
                        self.test_names.append(line.strip().split(":")[1].strip())
        except FileNotFoundError:
            pass  # Use default values if the file is not found